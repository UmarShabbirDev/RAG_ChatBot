import os
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain.schema import Document

def Context_Extractor(file_path):
    text_extension = ['.txt']
    pdf_extension = ['.pdf']
    docx_extension = ['.docx']
    text = ''
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension in text_extension:
        with open(file_path, 'r') as f:
            text = f.read()
    elif file_extension in pdf_extension:
        pdf_loader = PyPDFLoader(file_path)
        pages = pdf_loader.load()  # Correctly load the PDF pages
        for page in pages:
            text += page.page_content  # Access the content of each page correctly
            text += '\n'
    elif file_extension in docx_extension:
        doc = DocxDocument(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs]
        text = '\n'.join(text)
    else:
        print(f'File extension not supported for file: {file_path}')
        return None

    return Document(page_content=text, metadata={"source": file_path})

